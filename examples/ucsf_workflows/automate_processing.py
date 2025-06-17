#!/usr/bin/env python3
"""
UCSF Ac-225 iQID Processing Pipeline - automate_processing.py

This script handles the initial processing of raw iQID images, including:
- Metadata extraction from image headers
- Quantitative corrections (decay, calibration, unit conversion)
- ROI identification and extraction
- Data saving in structured format

Based on the original UCSF Ac-225 processing plan.
Integrates with unified UCSF data loader and quantitative corrections.

Author: Wookjin Choi <wookjin.choi@jefferson.edu>
Date: June 2025
"""

import os
import sys
import json
import logging
import numpy as np
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
import time
from datetime import datetime, timedelta
import cv2
from concurrent.futures import ProcessPoolExecutor, as_completed

# Add the package to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

# Import UCSF-specific modules
from ucsf_data_loader import UCSFDataMatcher
from quantitative_corrections import UCSFQuantitativeCorrections

try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not installed. Quantitative notes parsing will be limited.")
    Document = None

try:
    import tifffile
    from PIL import Image
except ImportError:
    print("Warning: tifffile or PIL not installed. Image loading capabilities limited.")
    tifffile = None
    Image = None


class QuantitativeCorrector:
    """
    Handles quantitative corrections for iQID data using UCSF calibration parameters.
    
    Uses the quantitative corrections from quantitative_notes.docx:
    - Geometric efficiency = 0.5
    - Spatial pileup calibration at 25 FPS = 23.8% loss
    - Frame-rate corrections
    - Spatial calibration (39.0 µm pixel size, 10-cm slice thickness)
    """
    
    def __init__(self, config: Dict[str, Any]):
        """Initialize with UCSF quantitative corrections."""
        self.config = config
        
        # Initialize UCSF quantitative corrections
        self.ucsf_corrections = UCSFQuantitativeCorrections()
        
        # Get calibration parameters
        self.calibration_params = self.ucsf_corrections.get_calibration_parameters()
        
        # Set up logging
        self.logger = logging.getLogger(__name__)
        
        self.logger.info("Initialized with UCSF quantitative corrections:")
        self.logger.info(f"  Geometric efficiency: {self.calibration_params['geometric_efficiency']}")
        self.logger.info(f"  Spatial pileup loss: {self.calibration_params['spatial_pileup_loss_25fps']*100:.1f}%")
        self.logger.info(f"  Combined correction factor: {self.calibration_params['combined_correction_factor']:.3f}")
    
    def apply_corrections(self, raw_data: np.ndarray, metadata: Dict[str, Any]) -> Tuple[np.ndarray, Dict[str, Any]]:
        """
        Apply quantitative corrections to raw iQID data.
        
        Args:
            raw_data: Raw iQID activity counts
            metadata: Image metadata including frame rate if available
            
        Returns:
            corrected_data: Quantitatively corrected activity
            correction_report: Detailed correction report
        """
        # Extract frame rate from metadata if available
        frame_rate = metadata.get('frame_rate', self.calibration_params['frame_rate_fps'])
        
        # Apply UCSF quantitative corrections
        corrected_data, correction_metadata = self.ucsf_corrections.apply_quantitative_corrections(
            raw_data, frame_rate
        )
        
        # Generate comprehensive correction report
        correction_report = self.ucsf_corrections.generate_correction_report(
            raw_data, corrected_data, correction_metadata
        )
        
        # Add processing metadata
        correction_report['processing_metadata'] = {
            'input_metadata': metadata,
            'processing_timestamp': time.strftime("%Y-%m-%d %H:%M:%S"),
            'correction_module': 'UCSF quantitative_corrections',
            'source_document': 'quantitative_notes.docx'
        }
        
        self.logger.info(f"Applied quantitative corrections:")
        self.logger.info(f"  Input range: {np.min(raw_data):.1f} - {np.max(raw_data):.1f}")
        self.logger.info(f"  Corrected range: {np.min(corrected_data):.1f} - {np.max(corrected_data):.1f}")
        self.logger.info(f"  Correction factor: {correction_metadata['combined_correction_factor']:.3f}")
        
        return corrected_data, correction_report
    
    def convert_to_ac225(self, alpha_activity: np.ndarray) -> np.ndarray:
        """Convert alpha particle activity to actual Ac-225 activity."""
        return self.ucsf_corrections.convert_to_ac225_activity(alpha_activity)
    
    def calculate_concentration(self, corrected_activity: np.ndarray, units: str = 'mBq_per_mm3') -> np.ndarray:
        """Calculate activity concentration with proper spatial calibration."""
        return self.ucsf_corrections.calculate_activity_concentration(corrected_activity, units)
    
    def get_spatial_calibration(self) -> Dict[str, float]:
        """Get spatial calibration parameters."""
        return self.ucsf_corrections.get_spatial_calibration()


class ROIExtractor:
    """Handles ROI identification and extraction."""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.params = config.get('automate_processing_params', {})
        
    def get_contours(self, image: np.ndarray) -> List[np.ndarray]:
        """Extract contours from image using threshold method."""
        method = self.params.get('roi_extraction_method', 'threshold')
        
        if method == 'threshold':
            threshold_value = self.params.get('threshold_value', 100)
            
            # Convert to uint8 if needed
            if image.dtype != np.uint8:
                image_norm = ((image - image.min()) / (image.max() - image.min()) * 255).astype(np.uint8)
            else:
                image_norm = image
                
            # Apply threshold
            _, binary = cv2.threshold(image_norm, threshold_value, 255, cv2.THRESH_BINARY)
            
            # Find contours
            contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            return contours
        else:
            logging.warning(f"ROI extraction method '{method}' not implemented")
            return []
    
    def get_ROIs(self, image: np.ndarray) -> List[Dict[str, Any]]:
        """Extract ROIs with area filtering."""
        contours = self.get_contours(image)
        min_area = self.params.get('minimum_roi_area_pixels', 50)
        
        rois = []
        for i, contour in enumerate(contours):
            area = cv2.contourArea(contour)
            if area >= min_area:
                # Create mask for this ROI
                mask = np.zeros(image.shape, dtype=np.uint8)
                cv2.fillPoly(mask, [contour], 255)
                
                # Calculate statistics
                roi_pixels = image[mask > 0]
                
                roi_data = {
                    'id': i,
                    'contour': contour,
                    'mask': mask,
                    'area_pixels': area,
                    'mean_activity': np.mean(roi_pixels),
                    'total_activity': np.sum(roi_pixels),
                    'max_activity': np.max(roi_pixels),
                    'centroid': tuple(map(int, np.mean(contour.reshape(-1, 2), axis=0)))
                }
                rois.append(roi_data)
                
        logging.info(f"Found {len(rois)} ROIs meeting criteria (min area: {min_area} pixels)")
        return rois


class QuantitativeNotesParser:
    """Parses quantitative notes document for correction factors."""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.correction_data = {}
        self._parse_document()
    
    def _parse_document(self):
        """Parse the quantitative notes document."""
        if not Document or not os.path.exists(self.docx_path):
            logging.warning(f"Cannot parse quantitative notes: {self.docx_path}")
            return
            
        try:
            doc = Document(self.docx_path)
            
            # Parse tables for correction factors
            for table in doc.tables:
                self._parse_table(table)
                
            # Parse text for additional parameters
            for paragraph in doc.paragraphs:
                self._parse_paragraph(paragraph.text)
                
            logging.info(f"Parsed quantitative notes with {len(self.correction_data)} entries")
            
        except Exception as e:
            logging.error(f"Error parsing quantitative notes: {e}")
    
    def _parse_table(self, table):
        """Parse a table from the document."""
        # Expected table structure: Animal ID | Injection Time | Activity | etc.
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        
        for row in table.rows[1:]:
            row_data = {}
            for i, cell in enumerate(row.cells):
                if i < len(headers):
                    row_data[headers[i]] = cell.text.strip()
            
            # Extract animal/sample ID
            animal_id = row_data.get('animal_id') or row_data.get('sample_id') or row_data.get('id')
            if animal_id:
                self.correction_data[animal_id] = row_data
    
    def _parse_paragraph(self, text: str):
        """Parse paragraph text for additional parameters."""
        # Look for key-value pairs like "Half-life: 10 days"
        import re
        
        patterns = {
            'half_life_days': r'half.?life[:\s]+(\d+\.?\d*)\s*days?',
            'injection_activity_mbq': r'injection[:\s]+(\d+\.?\d*)\s*mbq',
            'calibration_factor': r'calibration[:\s]+(\d+\.?\d*)'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text.lower())
            if match:
                self.correction_data[key] = float(match.group(1))
    
    def get_correction_data(self, sample_id: str) -> Dict[str, Any]:
        """Get correction data for a specific sample."""
        return self.correction_data.get(sample_id, {})


class AutomateProcessing:
    """Main processing class for raw iQID images using unified UCSF data loader."""
    
    def __init__(self, config_path: str = "unified_config.json"):
        self.config_path = config_path
        self.config = self._load_config()
        self._setup_logging()
        
        # Initialize UCSF data loader
        self.data_matcher = UCSFDataMatcher(self.config)
        
        # Initialize components
        self.corrector = QuantitativeCorrector(self.config)
        self.roi_extractor = ROIExtractor(self.config)
        
        # Setup directories
        self._setup_directories()
        
        # Log initialization
        logging.info(f"Initialized AutomateProcessing with {len(self.data_matcher.get_all_iqid_samples())} iQID samples")
        
    def _load_config(self) -> Dict[str, Any]:
        """Load configuration from JSON file."""
        try:
            with open(self.config_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.error(f"Config file not found: {self.config_path}")
            return self._get_default_config()
    
    def _get_default_config(self) -> Dict[str, Any]:
        """Get default unified configuration."""
        return {
            "experiment_id": "ucsf_ac225_iqid_processing",
            "base_path": "/home/wxc151/UCSF_Ac225_images",
            "paths": {
                "he_images": "H&E Images",
                "iqid_data": "iQID Data",
                "output_base": "output",
                "processed_data": "output/processed_data",
                "logs": "output/logs"
            },
            "sample_matching": {
                "he_pattern": r"([DM]\d+[A-Z]\d+)_([LR])",
                "iqid_pattern": r"([DM]\d+[A-Z]\d+)\(P\d+\)_([LR])",
                "case_sensitive": False
            },
            "processing": {
                "roi_extraction_method": "threshold",
                "threshold_value": 100,
                "minimum_roi_area_pixels": 50,
                "enable_parallel": True,
                "max_workers": 4
            },
            "logging": {
                "level": "INFO",
                "to_file": True
            }
        }
    
    def _setup_logging(self):
        """Setup logging configuration."""
        log_config = self.config.get('logging', {})
        log_level = getattr(logging, log_config.get('level', 'INFO'))
        
        logging.basicConfig(
            level=log_level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        
        if log_config.get('to_file', True):
            log_dir = self._get_full_path('logs')
            os.makedirs(log_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            log_file = os.path.join(log_dir, f"automate_processing_{timestamp}.log")
            
            file_handler = logging.FileHandler(log_file)
            file_handler.setLevel(log_level)
            file_handler.setFormatter(logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            ))
            
            logging.getLogger().addHandler(file_handler)
            
        logging.info(f"Logging setup complete for experiment: {self.config['experiment_id']}")
    
    def _get_full_path(self, path_key: str) -> str:
        """Convert path key to full path using unified config structure."""
        if path_key in self.config['paths']:
            relative_path = self.config['paths'][path_key]
            base_path = self.config['base_path']
            return os.path.join(base_path, relative_path)
        else:
            # Fallback for direct paths
            return path_key
    
    def _setup_directories(self):
        """Create necessary output directories."""
        dirs_to_create = [
            'output_base_dir',
            'processed_data_dir',
            'log_dir'
        ]
        
        for dir_key in dirs_to_create:
            if dir_key in self.config['paths']:
                dir_path = self._get_full_path(self.config['paths'][dir_key])
                os.makedirs(dir_path, exist_ok=True)
                logging.info(f"Created directory: {dir_path}")
    
    def find_raw_images(self) -> List[str]:
        """Find all raw iQID image files to process."""
        raw_dir = self._get_full_path(self.config['paths']['raw_image_dir'])
        
        if not os.path.exists(raw_dir):
            logging.warning(f"Raw image directory not found: {raw_dir}")
            return []
        
        # Common image file patterns
        patterns = ['*.tif', '*.tiff', '*.dcm', '*.nii', '*.nii.gz']
        image_files = []
        
        for pattern in patterns:
            files = glob.glob(os.path.join(raw_dir, '**', pattern), recursive=True)
            image_files.extend(files)
        
        logging.info(f"Found {len(image_files)} raw image files")
        return sorted(image_files)
    
    def extract_metadata(self, image_path: str) -> Dict[str, Any]:
        """Extract metadata from image file."""
        metadata = {
            'filename': os.path.basename(image_path),
            'filepath': image_path,
            'file_size_bytes': os.path.getsize(image_path)
        }
        
        try:
            if tifffile and image_path.lower().endswith(('.tif', '.tiff')):
                with tifffile.TiffFile(image_path) as tif:
                    if tif.pages:
                        page = tif.pages[0]
                        metadata.update({
                            'width': page.imagewidth,
                            'height': page.imagelength,
                            'dtype': str(page.dtype),
                            'shape': page.shape
                        })
                        
                        # Extract TIFF tags if available
                        if hasattr(page, 'tags'):
                            for tag in page.tags:
                                if tag.name in ['DateTime', 'XResolution', 'YResolution']:
                                    metadata[f'tiff_{tag.name.lower()}'] = tag.value
            
            elif Image and image_path.lower().endswith(('.jpg', '.png')):
                with Image.open(image_path) as img:
                    metadata.update({
                        'width': img.width,
                        'height': img.height,
                        'mode': img.mode
                    })
                    
                    # Extract EXIF data if available
                    if hasattr(img, '_getexif') and img._getexif():
                        exif = img._getexif()
                        if exif:
                            metadata['exif_datetime'] = exif.get(306)  # DateTime tag
            
        except Exception as e:
            logging.warning(f"Could not extract metadata from {image_path}: {e}")
        
        return metadata
    
    def load_image(self, image_path: str) -> np.ndarray:
        """Load image data as numpy array."""
        try:
            if tifffile and image_path.lower().endswith(('.tif', '.tiff')):
                return tifffile.imread(image_path)
            elif Image:
                img = Image.open(image_path)
                return np.array(img)
            else:
                # Fallback to OpenCV
                img = cv2.imread(image_path, cv2.IMREAD_UNCHANGED)
                if img is None:
                    raise ValueError(f"Could not load image: {image_path}")
                return img
                
        except Exception as e:
            logging.error(f"Failed to load image {image_path}: {e}")
            raise
    
    def process_single_image(self, image_path: str) -> Dict[str, Any]:
        """Process a single image file."""
        logging.info(f"Processing image: {os.path.basename(image_path)}")
        
        try:
            # Extract metadata
            metadata = self.extract_metadata(image_path)
            
            # Load image data
            image_data = self.load_image(image_path)
            
            # Get sample ID from filename (you may need to adjust this pattern)
            sample_id = self._extract_sample_id(image_path)
            
            # Get correction data for this sample
            correction_data = self.notes_parser.get_correction_data(sample_id)
            
            # Apply quantitative corrections
            corrected_data = image_data.astype(np.float32)
            
            # Apply decay correction if injection/acquisition times are available
            if 'injection_time' in correction_data and 'acquisition_time' in metadata:
                injection_time = self._parse_datetime(correction_data['injection_time'])
                acquisition_time = self._parse_datetime(metadata.get('acquisition_time', ''))
                
                if injection_time and acquisition_time:
                    corrected_data = self.corrector.apply_decay_correction(
                        corrected_data, injection_time, acquisition_time
                    )
            
            # Apply scanner calibration
            corrected_data = self.corrector.apply_scanner_calibration(corrected_data)
            
            # Apply unit conversion
            corrected_data = self.corrector.apply_unit_conversion(corrected_data)
            
            # Extract ROIs
            rois = self.roi_extractor.get_ROIs(corrected_data)
            
            # Prepare results
            results = {
                'sample_id': sample_id,
                'metadata': metadata,
                'correction_data': correction_data,
                'corrected_image_shape': corrected_data.shape,
                'corrected_image_stats': {
                    'min': float(np.min(corrected_data)),
                    'max': float(np.max(corrected_data)),
                    'mean': float(np.mean(corrected_data)),
                    'std': float(np.std(corrected_data))
                },
                'rois': [{
                    'id': roi['id'],
                    'area_pixels': roi['area_pixels'],
                    'mean_activity': roi['mean_activity'],
                    'total_activity': roi['total_activity'],
                    'max_activity': roi['max_activity'],
                    'centroid': roi['centroid']
                } for roi in rois],
                'processing_timestamp': datetime.now().isoformat()
            }
            
            # Save processed data
            self._save_processed_data(sample_id, corrected_data, rois, results)
            
            logging.info(f"Successfully processed {sample_id}: {len(rois)} ROIs found")
            return results
            
        except Exception as e:
            logging.error(f"Error processing {image_path}: {e}")
            return {'error': str(e), 'image_path': image_path}
    
    def _extract_sample_id(self, image_path: str) -> str:
        """Extract sample ID from filename."""
        filename = os.path.basename(image_path)
        
        # Common patterns for UCSF data
        patterns = [
            r'(D\d+M\d+[_\w]*)',  # D1M1_L pattern
            r'mouse(\d+)_day(\d+)',  # mouse01_day03 pattern
            r'(\w+)_(\d+)',  # general pattern
        ]
        
        for pattern in patterns:
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                return match.group(0)
        
        # Fallback: use filename without extension
        return os.path.splitext(filename)[0]
    
    def _parse_datetime(self, datetime_str: str) -> Optional[datetime]:
        """Parse datetime string to datetime object."""
        if not datetime_str:
            return None
            
        # Common datetime formats
        formats = [
            '%Y-%m-%d %H:%M:%S',
            '%Y/%m/%d %H:%M:%S',
            '%m/%d/%Y %H:%M:%S',
            '%Y-%m-%d',
            '%m/%d/%Y'
        ]
        
        for fmt in formats:
            try:
                return datetime.strptime(datetime_str, fmt)
            except ValueError:
                continue
        
        logging.warning(f"Could not parse datetime: {datetime_str}")
        return None
    
    def _save_processed_data(self, sample_id: str, corrected_data: np.ndarray, 
                           rois: List[Dict], results: Dict[str, Any]):
        """Save processed data to files."""
        output_dir = self._get_full_path(self.config['paths']['processed_data_dir'])
        sample_dir = os.path.join(output_dir, sample_id)
        os.makedirs(sample_dir, exist_ok=True)
        
        # Save corrected image data
        image_path = os.path.join(sample_dir, f"{sample_id}_corrected.npy")
        np.save(image_path, corrected_data)
        
        # Save ROI masks
        for roi in rois:
            mask_path = os.path.join(sample_dir, f"{sample_id}_roi_{roi['id']}_mask.npy")
            np.save(mask_path, roi['mask'])
        
        # Save metadata and results
        results_path = os.path.join(sample_dir, f"{sample_id}_results.json")
        with open(results_path, 'w') as f:
            json.dump(results, f, indent=2)
        
        logging.info(f"Saved processed data for {sample_id} to {sample_dir}")
    
    def run_processing(self, parallel: bool = None) -> Dict[str, Any]:
        """Run the complete processing pipeline."""
        if parallel is None:
            parallel = self.config.get('parallelization', {}).get('enabled', True)
        
        image_files = self.find_raw_images()
        
        if not image_files:
            logging.warning("No image files found to process")
            return {'status': 'no_files', 'results': []}
        
        logging.info(f"Starting processing of {len(image_files)} images (parallel: {parallel})")
        
        start_time = time.time()
        results = []
        
        if parallel and len(image_files) > 1:
            max_workers = self.config.get('parallelization', {}).get('max_workers_processing', 4)
            
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {
                    executor.submit(self.process_single_image, image_file): image_file 
                    for image_file in image_files
                }
                
                for future in as_completed(future_to_file):
                    image_file = future_to_file[future]
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        logging.error(f"Processing failed for {image_file}: {e}")
                        results.append({'error': str(e), 'image_path': image_file})
        else:
            # Sequential processing
            for image_file in image_files:
                result = self.process_single_image(image_file)
                results.append(result)
        
        end_time = time.time()
        processing_time = end_time - start_time
        
        # Summary statistics
        successful = len([r for r in results if 'error' not in r])
        failed = len(results) - successful
        
        summary = {
            'status': 'completed',
            'total_images': len(image_files),
            'successful': successful,
            'failed': failed,
            'processing_time_seconds': processing_time,
            'results': results,
            'experiment_id': self.config['experiment_id'],
            'completion_timestamp': datetime.now().isoformat()
        }
        
        # Save summary
        summary_path = os.path.join(
            self._get_full_path(self.config['paths']['output_base_dir']),
            'processing_summary.json'
        )
        
        with open(summary_path, 'w') as f:
            json.dump(summary, f, indent=2)
        
        logging.info(f"Processing completed: {successful}/{len(image_files)} successful "
                    f"in {processing_time:.1f} seconds")
        
        return summary


def main():
    """Main execution function."""
    print("🔬 UCSF Ac-225 iQID Processing - automate_processing.py")
    print("=" * 60)
    
    # Check for config file
    config_path = "config.json"
    if len(sys.argv) > 1:
        config_path = sys.argv[1]
    
    if not os.path.exists(config_path):
        print(f"❌ Config file not found: {config_path}")
        print("Creating default config file...")
        
        # Create default config
        processor = AutomateProcessing("nonexistent.json")  # Will use defaults
        with open(config_path, 'w') as f:
            json.dump(processor.config, f, indent=2)
        
        print(f"✅ Default config created: {config_path}")
        print("Please review and modify the configuration as needed, then run again.")
        return
    
    try:
        # Initialize processor
        processor = AutomateProcessing(config_path)
        
        # Run processing
        results = processor.run_processing()
        
        # Display summary
        print(f"\n✅ Processing completed!")
        print(f"📊 Results: {results['successful']}/{results['total_images']} successful")
        print(f"⏱️  Processing time: {results['processing_time_seconds']:.1f} seconds")
        
        if results['failed'] > 0:
            print(f"⚠️  {results['failed']} files failed to process")
        
        print(f"📁 Outputs saved to: {processor._get_full_path(processor.config['paths']['processed_data_dir'])}")
        
    except Exception as e:
        print(f"❌ Error during processing: {e}")
        logging.error(f"Fatal error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
